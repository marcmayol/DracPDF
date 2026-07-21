"""Criterio de aceptación funcional de la Fase 7 (conversiones), sin UI.

Sobre fixtures de prueba: convierte un PDF a docx (texto + tabla), a HTML y
Markdown (texto + títulos), un docx a PDF (texto seleccionable + tabla), y detecta
un PDF escaneado. Exit 0 si todo pasa; muestra además la versión de PyMuPDF.

Word→PDF usa Qt y necesita fuentes, así que se fuerza la plataforma NATIVA (no
offscreen).

Uso:
    uv run python scripts/verificar_conversiones.py
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.pop("QT_QPA_PLATFORM", None)  # plataforma nativa (con fuentes)
sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # consola cp1252 → utf-8

import docx  # noqa: E402
import fitz  # noqa: E402
from PySide6.QtWidgets import QApplication  # noqa: E402

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from tests.adapters.generar_fixtures import (  # noqa: E402
    generar_pdf_escaneado,
    generar_pdf_titulos_tabla,
)
from tests.adapters.generar_fixtures_docx import generar_docx_prueba  # noqa: E402

from lectorpdf.adapters.pymupdf.conversor import ConversorFitz  # noqa: E402
from lectorpdf.adapters.pymupdf.document_repository import (  # noqa: E402
    PyMuPDFDocumentRepository,
)
from lectorpdf.adapters.pymupdf.registro import RegistroDocumentos  # noqa: E402
from lectorpdf.adapters.qt.conversor_word import ConversorWordQt  # noqa: E402
from lectorpdf.core.domain.conversion import A4  # noqa: E402

_resultados: list[tuple[str, bool]] = []


def _check(nombre: str, condicion: bool) -> None:
    _resultados.append((nombre, condicion))


def main() -> int:
    QApplication.instance() or QApplication([])
    tmp = Path(tempfile.mkdtemp())

    registro = RegistroDocumentos()
    pdf = generar_pdf_titulos_tabla(tmp / "origen.pdf")
    doc = PyMuPDFDocumentRepository(registro).abrir(pdf)
    conv = ConversorFitz(registro)

    # 1. PDF → Word (texto + tabla).
    a_docx = tmp / "salida.docx"
    conv.a_word(doc.id, a_docx)
    d = docx.Document(str(a_docx))
    celdas = [c.text for t in d.tables for f in t.rows for c in f.cells]
    texto_docx = "\n".join(p.text for p in d.paragraphs)
    tabla_ok = len(d.tables) >= 1 and any("Ingresos" in c for c in celdas)
    _check("PDF→Word: texto", "Informe de prueba" in texto_docx)
    _check("PDF→Word: al menos una tabla", tabla_ok)

    # 2. PDF → HTML (texto + títulos).
    a_html = tmp / "salida.html"
    conv.a_html(doc.id, a_html)
    html = a_html.read_text(encoding="utf-8")
    _check("PDF→HTML: texto y títulos", "Informe de prueba" in html and "segunda" in html)

    # 3. PDF → Markdown (texto + títulos).
    a_md = tmp / "salida.md"
    conv.a_markdown(doc.id, a_md)
    md = a_md.read_text(encoding="utf-8")
    titulo = any(ln.startswith("#") and "Informe de prueba" in ln for ln in md.splitlines())
    _check("PDF→Markdown: título con # y tabla", titulo and "| --- |" in md)

    # 4. Word → PDF (texto seleccionable + tabla).
    docx_fixture = generar_docx_prueba(tmp / "prueba.docx")
    a_pdf = tmp / "desde_word.pdf"
    ConversorWordQt().a_pdf(docx_fixture, a_pdf, A4)
    pdf_doc = fitz.open(a_pdf)
    texto_pdf = "\n".join(pdf_doc[i].get_text("text") for i in range(pdf_doc.page_count))
    fuentes = sum(len(pdf_doc.get_page_fonts(i)) for i in range(pdf_doc.page_count))
    pdf_doc.close()
    seleccionable = fuentes > 0 and "Contrato" in texto_pdf
    tabla_pdf = "Renta mensual" in texto_pdf and "750 EUR" in texto_pdf
    _check("Word→PDF: texto seleccionable (con fuentes)", seleccionable)
    _check("Word→PDF: tabla renderizada", tabla_pdf)

    # 5. Detección de PDF escaneado.
    esc = generar_pdf_escaneado(tmp / "escaneado.pdf")
    doc_esc = PyMuPDFDocumentRepository(registro).abrir(esc)
    _check("PDF escaneado detectado (aviso)", conv.es_escaneado(doc_esc.id) is True)

    print("-" * 60)
    ok = True
    for nombre, cond in _resultados:
        print(f"  [{'OK' if cond else 'FALLO'}] {nombre}")
        ok = ok and cond
    print("-" * 60)
    print("PyMuPDF:", fitz.VersionBind, "(sin downgrade; el proyecto fija >=1.24)")
    print("RESULTADO:", "OK" if ok else "FALLO")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())

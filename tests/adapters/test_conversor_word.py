"""Test de integración PDF → Word (pdf2docx) del adaptador ConversorFitz."""

from __future__ import annotations

from pathlib import Path

import docx

from lectorpdf.adapters.pymupdf.conversor import ConversorFitz
from lectorpdf.adapters.pymupdf.document_repository import PyMuPDFDocumentRepository
from lectorpdf.adapters.pymupdf.registro import RegistroDocumentos


def _abrir(ruta: Path) -> tuple[ConversorFitz, str]:
    registro = RegistroDocumentos()
    documento = PyMuPDFDocumentRepository(registro).abrir(ruta)
    return ConversorFitz(registro), documento.id


def test_a_word_genera_docx_con_texto_y_tabla(pdf_titulos_tabla: Path, tmp_path: Path) -> None:
    conversor, doc_id = _abrir(pdf_titulos_tabla)
    destino = tmp_path / "salida.docx"

    conversor.a_word(doc_id, destino)

    assert destino.is_file()
    documento = docx.Document(str(destino))
    texto = "\n".join(p.text for p in documento.paragraphs)
    texto += "\n".join(
        celda.text for tabla in documento.tables for fila in tabla.rows for celda in fila.cells
    )
    assert "Informe de prueba" in texto  # título
    assert len(documento.tables) >= 1  # al menos una tabla
    # La tabla conserva las celdas.
    celdas = [c.text for t in documento.tables for f in t.rows for c in f.cells]
    assert any("Ingresos" in c for c in celdas)


def test_a_word_reporta_progreso(pdf_titulos_tabla: Path, tmp_path: Path) -> None:
    conversor, doc_id = _abrir(pdf_titulos_tabla)
    hechos: list[tuple[int, int]] = []

    conversor.a_word(
        doc_id, tmp_path / "s.docx", progreso=lambda h, t: hechos.append((h, t))
    )

    assert hechos and hechos[-1] == (1, 1)

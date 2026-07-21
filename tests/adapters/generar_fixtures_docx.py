"""Genera un .docx de prueba para las conversiones Word→PDF.

Uso a mano:
    uv run python tests/adapters/generar_fixtures_docx.py
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Inches


def _png_pequeno() -> Path:
    """Un PNG pequeño (para add_picture); python-docx necesita un fichero."""
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 64, 64), False)
    pix.set_rect(pix.irect, (60, 120, 200))
    fd, nombre = tempfile.mkstemp(suffix=".png")
    os.close(fd)  # en Windows, el fd abierto bloquearía el fichero
    pix.save(nombre)
    return Path(nombre)


def generar_docx_prueba(destino: Path) -> Path:
    """DOCX con título, subtítulo, párrafo con negrita, lista, tabla e imagen."""
    doc = Document()
    doc.add_heading("Contrato de arrendamiento", level=1)
    parrafo = doc.add_paragraph("Este documento incluye ")
    parrafo.add_run("texto en negrita").bold = True
    parrafo.add_run(" y contenido variado.")

    doc.add_heading("Cláusulas", level=2)
    doc.add_paragraph("Primera cláusula del contrato.", style="List Bullet")
    doc.add_paragraph("Segunda cláusula del contrato.", style="List Bullet")

    tabla = doc.add_table(rows=2, cols=2)
    tabla.style = "Table Grid"
    tabla.cell(0, 0).text = "Concepto"
    tabla.cell(0, 1).text = "Importe"
    tabla.cell(1, 0).text = "Renta mensual"
    tabla.cell(1, 1).text = "750 EUR"

    imagen = _png_pequeno()
    try:
        doc.add_picture(str(imagen), width=Inches(1.0))
    finally:
        imagen.unlink(missing_ok=True)

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return destino


if __name__ == "__main__":
    ruta = generar_docx_prueba(Path("build") / "fixtures" / "prueba.docx")
    print("Generado:", ruta)
